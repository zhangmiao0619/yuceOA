#!/usr/bin/env python3
"""Rebuild chapter 4 — correct order"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PATH = '/Volumes/YUCE/宇测OA 系统/oa-system/docs/OA/项目管理模块使用手册.docx'
doc = Document(PATH)
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
body = doc.element.body

def ptext(el):
    t = []
    for tx in el.iter(f'{{{W}}}t'):
        if tx.text:
            t.append(tx.text)
    return ''.join(t).strip()

# Find h3 and h5 positions
children = list(body)
h3_idx = h5_idx = None
for i, c in enumerate(children):
    tag = c.tag.split('}')[-1]
    if tag == 'p':
        txt = ptext(c)
        if txt == '3. 创建项目':
            h3_idx = i
        elif txt.startswith('5. 项目'):
            h5_idx = i

print(f'h3={h3_idx}, h5={h5_idx}')

# Remove everything between h3 and h5
for i in range(h5_idx - 1, h3_idx, -1):
    body.remove(children[i])

# Find current h3 element reference
h3_el = None
h5_el = None
for c in body:
    tag = c.tag.split('}')[-1]
    if tag == 'p':
        txt = ptext(c)
        if txt == '3. 创建项目':
            h3_el = c
        elif txt.startswith('5. 项目'):
            h5_el = c

def make_style(run, bold=False, size=11, font='微软雅黑'):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)

def ah(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        make_style(r)
    return h._element

def ap(text, bold=False, size=11, bullet=False, sa=6):
    if bullet:
        p = doc.add_paragraph(style='List Bullet')
        p.text = ''
        p.add_run(text)
    else:
        p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(sa)
    for r in p.runs:
        make_style(r, bold=bold, size=size)
    return p._element

def at(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
            r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, v in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(v)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(10)
    return t._tbl

# Build elements in order
els = []

els.append(ah('4. 项目详情页', 1))
els.append(ap('点击项目名称进入详情页，包含两个标签页。'))

els.append(ah('4.1 概览标签', 2))
els.append(ap('展示项目基本信息：名称、状态、负责人、起止日期、客户简称、工作量与计量单位。顶部有"编辑"按钮，可在弹出的对话框中修改以下字段：'))
els.append(at(['字段', '说明'], [
    ['项目名称', '必填'],
    ['开始日期', '修改项目开始日期'],
    ['结束日期', '修改项目结束日期，不能早于开始日期'],
    ['工作量', '修改工作量数值'],
    ['计量单位', '天/平方米/平方千米/米/千米/亩/栋/宗/块'],
]))
els.append(ap('保存后项目管理列表数据同步更新。'))

els.append(ah('4.2 子任务标签', 2))
els.append(ah('4.2.1 创建子任务', 3))
els.append(ap('① 切换到"子任务"标签，点击"新建子任务"按钮。'))
els.append(ap('② 填写子任务信息：'))
els.append(at(['字段', '必填', '说明'], [
    ['子任务名称', '是', '任务标题'],
    ['子任务描述', '否', '多行文本，任务要求说明'],
    ['任务级别', '否', 'P0 / P1 / P2 / P3，默认 P2'],
    ['负责人', '是', '从员工列表选择指派人'],
    ['计划开始时间', '否', '日期选择'],
    ['预计结束时间', '否', '不能早于计划开始时间'],
    ['工作量', '否', '数字，支持小数'],
    ['计量单位', '否', '天/平方米/平方千米/米/千米/亩/栋/宗/块'],
    ['子任务备注', '否', '最长 200 字'],
]))
els.append(ap('③ 保存后创建成功，负责人收到通知。'))

els.append(ah('4.2.2 任务状态流转', 3))
els.append(at(['当前状态', '下一状态', '操作'], [
    ['未指派', '待接收', '编辑任务，指定负责人'],
    ['待接收', '进行中', '负责人点"接收任务"'],
    ['进行中', '待审核', '负责人点"提交完成"'],
    ['待审核', '已完成', '审核人点"审核通过"'],
    ['待审核', '进行中', '审核人点"驳回"'],
    ['已完成', '—', '终态，不可再变'],
]))
els.append(ap(''))
els.append(ap('编辑：修改任务信息（标题、描述、优先级、负责人等）', bullet=True))
els.append(ap('删除：移除未开始的任务', bullet=True))
els.append(ap('点击任务标题可查看完整详情', bullet=True))

print(f'Created {len(els)} elements')

# Insert each AFTER h3_el's last sibling, in order
# Insert el before h5, in order — each time h5 shifts right
for el in els:
    h5_el.addprevious(el)

doc.save(PATH)
print('Saved — chapter 4 updated in correct order')
